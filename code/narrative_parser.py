import os
import json
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any, Optional, Tuple


class NarrativeParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.full_text = self._extract_full_text()
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.table_data = self._extract_table_data_robust()

    def _extract_full_text(self) -> str:
        full_text = []

        for para in self.doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        return '\n'.join(full_text)

    def _clean_key(self, text: str) -> str:
        text = re.sub(r'\*+', '', text)
        text = re.sub(r'\s+', ' ', text)
        text = text.strip(':').strip()
        return text

    def _is_valid_label(self, text: str) -> bool:
        if not text or len(text) < 3:
            return False

        if not re.search(r'[a-zA-Z]', text):
            return False

        if len(text) > 200:
            return False
        
        label_keywords = [
            'project', 'title', 'budget', 'coordinator', 'name', 'date',
            'location', 'program', 'year', 'student', 'id', 'implementation',
            'beneficiaries', 'results', 'objectives', 'goals', 'team', 'participant'
        ]

        text_lower = text.lower()
        has_keyword = any(kw in text_lower for kw in label_keywords)

        ends_properly = text.endswith(':') or text.endswith('?') or text[0].isupper()

        return has_keyword or ends_properly

    def _is_valid_value(self, text: str, key: str) -> bool:
        if not text:
            return False

        if text.endswith('?'):
            return False

        key_lower = key.lower()

        if 'budget' in key_lower:
            return bool(re.search(r'\d', text))

        if 'date' in key_lower:
            return bool(re.search(r'(?:\d{4}|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*)', text,
                                  re.IGNORECASE))

        if self._is_valid_label(text) and len(text) < 100:
            return False

        return True

    def _extract_table_data_robust(self) -> Dict[str, str]:
        """Extract data from tables with robust validation."""
        table_data = {}

        for table in self.doc.tables:
            num_cols = len(table.columns)

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                if not any(cells):
                    continue

                if num_cols >= 2:
                    key = self._clean_key(cells[0])
                    value = cells[1].strip()

                    if (self._is_valid_label(key) and
                            value and
                            self._is_valid_value(value, key) and
                            key != value):

                        if key not in table_data or len(value) > len(table_data[key]):
                            table_data[key] = value

        return table_data

    def _search_in_text(self, patterns: List[str], context_lines: int = 3) -> str:
        for pattern in patterns:
            regex = rf'{re.escape(pattern)}\s*:?\s*(.+?)(?:\n|$)'
            match = re.search(regex, self.full_text, re.IGNORECASE)

            if match:
                value = match.group(1).strip()
                if value and len(value) < 200 and not value.endswith('?'):
                    return value

        return ''

    def extract_project_title(self) -> str:
        for key, value in self.table_data.items():
            if re.search(r'\bproject\s+title\b|\btitle\b', key, re.IGNORECASE):
                if not self._is_valid_label(value):
                    return value
                if not re.search(r'budget|coordinator|program|year|location|date', value, re.IGNORECASE):
                    return value

        patterns = [
            'Project Title',
            'Project Name',
            'Title of Project',
            'Name of Project'
        ]
        result = self._search_in_text(patterns)

        if result and not re.search(r'budget|coordinator|program\s+id', result, re.IGNORECASE):
            return result

        return ''

    def extract_project_budget(self) -> str:
        for key, value in self.table_data.items():
            if re.search(r'\bbudget\b', key, re.IGNORECASE):
                if re.search(r'\d', value):
                    if not re.search(r'coordinator|program|year|location|project title', value, re.IGNORECASE):
                        return value

        currency_patterns = [
            r'(?:budget|funding|grant).*?(\d[\d,\s]*\d*\s*(?:KGZ|KGS|USD|EUR|som|dollars?))',
            r'(\d[\d,\s]+\s*(?:KGZ|KGS|USD|EUR|som))',
        ]

        for pattern in currency_patterns:
            match = re.search(pattern, self.full_text, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        return ''

    def extract_coordinator(self) -> str:
        for key, value in self.table_data.items():
            if re.search(r'\bcoordinator\b|\bleader\b|\bmanager\b', key, re.IGNORECASE):
                if not re.search(r'program|year|student|date|location|budget', value, re.IGNORECASE):
                    if len(value) < 100:
                        return value

        patterns = [
            'Project Coordinator',
            'Coordinator',
            'Project Leader',
            'Team Leader'
        ]
        return self._search_in_text(patterns)

    def extract_dates(self) -> Dict[str, str]:
        dates = {}
        date_patterns = [
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}?,?\s*\d{4}',
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}',
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{4}-\d{2}-\d{2}'
        ]

        for key, value in self.table_data.items():
            if re.search(r'implementation|duration|timeline|project\s+dates?', key, re.IGNORECASE):
                if any(re.search(pattern, value, re.IGNORECASE) for pattern in date_patterns):
                    if not re.search(r'location|coordinator|budget|program', value, re.IGNORECASE):
                        dates['overall'] = value
                        break

        search_text = self.full_text

        online_match = re.search(
            r'((?:September|January|February|March|April|May|June|July|August|October|November|December)\s+\d{4}\s*[-–]\s*(?:September|January|February|March|April|May|June|July|August|October|November|December)\s+\d{4})\s*online',
            search_text,
            re.IGNORECASE
        )
        if online_match:
            dates['online'] = online_match.group(1).strip()

        camp_match = re.search(
            r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}[-–]\d{1,2},?\s+\d{4})\s*(?:offline|camp)',
            search_text,
            re.IGNORECASE
        )
        if camp_match:
            dates['camp'] = camp_match.group(1).strip()

        return dates

    def extract_location(self) -> List[str]:
        locations = []

        location_text = ''
        for key, value in self.table_data.items():
            if re.search(r'\blocation\b|\bvenue\b', key, re.IGNORECASE):
                if not re.search(r'coordinator|budget|program|year|student', value, re.IGNORECASE):
                    location_text = value
                    break

        if not location_text:
            location_text = self.full_text[:2000]

        if re.search(r'\bonline\b', location_text, re.IGNORECASE):
            locations.append('online')

        cities = ['Bishkek', 'Naryn', 'Osh', 'Jalal-Abad', 'Batken', 'Karakol', 'Tokmok', 'Kara-Balta']
        for city in cities:
            if re.search(rf'\b{city}\b', location_text, re.IGNORECASE):
                if city not in locations:
                    locations.append(city)

        school_patterns = [
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:High\s+)?School',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:University|College)',
        ]

        for pattern in school_patterns:
            for match in re.finditer(pattern, location_text):
                institution = match.group(0).strip()
                if institution not in locations and len(institution) < 100:
                    locations.append(institution)

        return locations

    def extract_beneficiaries(self) -> Dict[str, int]:
        beneficiaries = {}

        ben_section_match = re.search(
            r'(?:How many beneficiaries|Number of beneficiaries|beneficiaries).*?(?=\n\n[A-Z]|\Z)',
            self.full_text,
            re.IGNORECASE | re.DOTALL
        )

        search_text = ben_section_match.group(0) if ben_section_match else self.full_text[:3000]

        patterns = [
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*mentors?\b', 'mentors'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*mentees?\b', 'mentees'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*staff\b', 'staff'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*camp\s*counselors?\b', 'camp_counselors'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*participants?\b', 'participants'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*students?\b', 'students'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*volunteers?\b', 'volunteers'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*trainers?\b', 'trainers'),
            (r'(?:^|\n)\s*(\d+)\s*[-–:]*\s*facilitators?\b', 'facilitators'),
        ]

        for pattern, key in patterns:
            match = re.search(pattern, search_text, re.IGNORECASE | re.MULTILINE)
            if match:
                num = int(match.group(1))
                if 0 < num < 10000:
                    beneficiaries[key] = num

        return beneficiaries

    def extract_list_section(self, section_markers: List[str], max_distance: int = 1000) -> List[str]:
        items = []

        section_text = None
        for marker in section_markers:
            pattern = rf'{re.escape(marker)}(.{{0,{max_distance}}}?)(?:\n\n[A-Z*]|\Z)'
            match = re.search(pattern, self.full_text, re.IGNORECASE | re.DOTALL)

            if match:
                section_text = match.group(1)
                break

        if not section_text:
            return items

        numbered = re.findall(r'(?:^|\n)\s*(\d+)\.\s+([^\n]+)', section_text, re.MULTILINE)
        if numbered:
            items.extend([item[1].strip() for item in numbered if item[1].strip() and len(item[1].strip()) > 5])

        bullets = re.findall(r'(?:^|\n)\s*[●•▪▫■□\-\*]\s+([^\n]+)', section_text, re.MULTILINE)
        if bullets and not items:
            items.extend([item.strip() for item in bullets if item.strip() and len(item.strip()) > 5])

        dashes = re.findall(r'(?:^|\n)\s*[-–—]\s+([^\n]+)', section_text, re.MULTILINE)
        if dashes and not items:
            items.extend([item.strip() for item in dashes if item.strip() and len(item.strip()) > 5])

        return items[:50]

    def extract_results(self) -> List[str]:
        markers = [
            'What were the project results?',
            'Project results',
            'Results:',
            'Results',
            'Outcomes',
            'Achievements',
            'What tangible or intangible results'
        ]
        return self.extract_list_section(markers, max_distance=2000)

    def extract_activities(self) -> List[str]:
        markers = [
            'Please describe project activities',
            'Project activities',
            'Activities',
            'WHAT was done',
            'Implementation activities'
        ]
        return self.extract_list_section(markers, max_distance=3000)

    def extract_team(self) -> List[str]:
        markers = [
            'List your team members',
            'Team members',
            'Team:',
            'How many members'
        ]
        return self.extract_list_section(markers, max_distance=1500)

    def validate_and_clean(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if data['project_budget']:
            budget = data['project_budget']
            budget = re.sub(r'\s+', ' ', budget).strip()
            data['project_budget'] = budget

        if data['project_title']:
            title = data['project_title']
            title = re.sub(r'^[\'"\s]+', '', title)
            title = re.sub(r'[.\'"\s]+$', '', title)
            if re.search(r'budget|coordinator|program|year of entry|student id', title, re.IGNORECASE):
                title = ''
            data['project_title'] = title

        if data['project_coordinator']:
            coord = data['project_coordinator']
            if re.search(r'program|year|entry|student|date|location|budget', coord, re.IGNORECASE):
                data['project_coordinator'] = ''

        return data

    def parse(self) -> Optional[Dict[str, Any]]:
        """Parse the entire document and return structured data."""
        data = {
            'project_title': self.extract_project_title(),
            'project_budget': self.extract_project_budget(),
            'project_coordinator': self.extract_coordinator(),
            'dates': self.extract_dates(),
            'location': self.extract_location(),
            'beneficiaries': self.extract_beneficiaries(),
            'results': self.extract_results(),
            'activities': self.extract_activities(),
            'team': self.extract_team(),
            'raw_text_excerpt': self.full_text[:1000] + '...' if len(self.full_text) > 1000 else self.full_text,
            'source_file': os.path.basename(self.docx_path)
        }

        data = self.validate_and_clean(data)

        return data


def process_all_reports(input_dir: str, output_dir: str):
    """Process all .docx files in the input directory."""
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    output_path.mkdir(parents=True, exist_ok=True)

    docx_files = list(input_path.glob('*.docx'))

    docx_files = [f for f in docx_files if not f.name.startswith('~$')]


    successful = 0
    failed = 0
    skipped = 0
    warnings = []

    for docx_file in docx_files:
        try:

            parser = NarrativeParser(str(docx_file))
            parsed_data = parser.parse()

            output_filename = docx_file.stem + '.json'
            output_file = output_path / output_filename

            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(parsed_data, f, indent=2, ensure_ascii=False)

            title = parsed_data['project_title'][:60] if parsed_data['project_title'] else 'NOT FOUND'
            budget = parsed_data['project_budget'] if parsed_data['project_budget'] else 'NOT FOUND'
            coord = parsed_data['project_coordinator'][:40]

            successful += 1

        except Exception as e:
            import traceback
            failed += 1

def main():
    input_directory = '../data/narratives'
    output_directory = '../data/narratives_parsed'

    process_all_reports(input_directory, output_directory)

if __name__ == '__main__':
    main()